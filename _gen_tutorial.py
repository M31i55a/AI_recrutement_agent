"""Script to generate the project tutorial .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_code_block(doc, code: str):
    """Add a shaded code block paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # Add shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    # Add border
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)
    return p


def add_info_box(doc, text: str, color="E8F4FD"):
    """Add a light-blue info box."""
    p = doc.add_paragraph()
    run = p.add_run("💡  " + text)
    run.font.size = Pt(10)
    run.font.italic = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def build_tutorial():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────
    title = doc.add_heading("AI Recruiter Agency", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_color(title, 0, 102, 204)

    subtitle = doc.add_paragraph("A Beginner's Complete Tutorial")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    desc = doc.add_paragraph(
        "This tutorial walks you through every part of the AI Recruiter Agency project — "
        "from the database all the way to the AI agents — so you understand exactly how "
        "each piece was built and why it works the way it does."
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. Project Overview
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("1. What Is This Project?", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "The AI Recruiter Agency is a web application that acts like a smart HR assistant. "
        "When a job candidate uploads their PDF resume, the system:"
    )
    for step in [
        "Reads and extracts all the text from the PDF.",
        "Understands the candidate's skills, experience, and education.",
        "Searches a database of job listings to find the best matches.",
        "Screens the candidate against job requirements.",
        "Produces a final recommendation with clear next steps.",
    ]:
        p = doc.add_paragraph(step, style="List Number")
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph()
    add_info_box(
        doc,
        "All the 'understanding' is done by a local AI model called Llama 3.2, "
        "running through a tool called Ollama — no internet or paid API is required!",
    )

    # ── Tech stack ────────────────────────────────────────────────
    doc.add_paragraph()
    h2 = doc.add_heading("Technology Stack", 2)
    set_heading_color(h2, 0, 128, 0)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Technology"
    hdr[1].text = "What It Does in This Project"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0066CC")
        tcPr.append(shd)

    rows_data = [
        ("Python 3.x", "The programming language used for the entire project."),
        ("Streamlit", "Builds the interactive web interface (no HTML/CSS needed)."),
        ("Ollama + Llama 3.2", "The local AI model that understands and generates text."),
        ("OpenAI Python SDK", "Used as a client library to talk to Ollama's API."),
        ("SQLite", "A lightweight file-based database that stores job listings."),
        ("pdfminer.six", "Extracts raw text from uploaded PDF resumes."),
        ("python-docx", "Used to create this very tutorial document!"),
    ]
    for tech, desc in rows_data:
        row = table.add_row().cells
        row[0].text = tech
        row[1].text = desc
        row[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. Project Structure
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("2. Project Folder Structure", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "Here is how all the files are organized. Understanding where everything lives "
        "makes it much easier to follow the code."
    )
    doc.add_paragraph()

    structure = (
        "ollama/                    ← Root project folder\n"
        "│\n"
        "├── app.py                 ← The web application (entry point)\n"
        "├── requirements.txt       ← List of Python packages to install\n"
        "│\n"
        "├── agents/                ← All AI agents live here\n"
        "│   ├── __init__.py\n"
        "│   ├── base_agent.py      ← Shared base class for all agents\n"
        "│   ├── orchestrator.py    ← Manages the entire workflow\n"
        "│   ├── extractor_agent.py ← Reads PDF and extracts data\n"
        "│   ├── analyzer_agent.py  ← Analyzes skills and experience\n"
        "│   ├── matcher_agent.py   ← Matches candidate to jobs\n"
        "│   ├── screener_agent.py  ← Screens the candidate\n"
        "│   └── recommender_agent.py ← Final recommendation\n"
        "│\n"
        "├── db/                    ← Database layer\n"
        "│   ├── database.py        ← Python class to interact with DB\n"
        "│   ├── schema.sql         ← Defines the database table structure\n"
        "│   └── seed_jobs.py       ← Fills the DB with sample job listings\n"
        "│\n"
        "├── utils/                 ← Helper utilities\n"
        "│   ├── logger.py          ← Sets up logging to files and console\n"
        "│   └── exceptions.py      ← Custom error types\n"
        "│\n"
        "├── uploads/               ← Uploaded resumes are saved here\n"
        "├── results/               ← Analysis output text files\n"
        "└── logs/                  ← Application log files"
    )
    add_code_block(doc, structure)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. The Database
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("3. The Database Layer", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "The database is the project's memory — it stores all the available job listings "
        "so that the AI can search through them when matching a candidate."
    )

    # 3.1
    h2 = doc.add_heading("3.1  Why SQLite?", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "SQLite was chosen because it requires zero setup: it is just a single file "
        "(jobs.sqlite) that Python can read and write directly. There is no separate "
        "database server to install or configure — perfect for a self-contained project."
    )

    # 3.2
    h2 = doc.add_heading("3.2  The Schema — db/schema.sql", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "The schema file defines the shape of the database — think of it like designing "
        "a spreadsheet template before you fill it with data."
    )
    add_code_block(
        doc,
        "CREATE TABLE IF NOT EXISTS jobs (\n"
        "    id               INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "    title            TEXT NOT NULL,\n"
        "    company          TEXT NOT NULL,\n"
        "    location         TEXT NOT NULL,\n"
        "    type             TEXT NOT NULL,   -- e.g. Full-time, Part-time\n"
        "    experience_level TEXT NOT NULL,   -- Junior / Mid-level / Senior\n"
        "    salary_range     TEXT,\n"
        "    description      TEXT NOT NULL,\n"
        "    requirements     TEXT NOT NULL,   -- stored as a JSON list\n"
        "    benefits         TEXT,            -- stored as a JSON list\n"
        "    created_at       DATETIME DEFAULT CURRENT_TIMESTAMP,\n"
        "    updated_at       DATETIME DEFAULT CURRENT_TIMESTAMP\n"
        ");",
    )

    doc.add_paragraph()
    doc.add_paragraph("Key design decisions explained:")
    bullets = [
        ("id INTEGER PRIMARY KEY AUTOINCREMENT",
         "Every job gets a unique number automatically — you never have to assign IDs yourself."),
        ("TEXT NOT NULL",
         "These columns are required — you cannot add a job without them."),
        ("requirements TEXT (JSON)",
         "Instead of creating a separate table for requirements, we store them as a JSON "
         "list inside a text column. Example: '[\"Python\", \"React\", \"AWS\"]'. This is "
         "simpler for a project of this scale."),
        ("created_at / updated_at",
         "Timestamps are added automatically so you can sort jobs by recency."),
    ]
    for term, explanation in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run_bold = p.add_run(term + ": ")
        run_bold.font.bold = True
        run_bold.font.name = "Courier New"
        p.add_run(explanation)

    # 3.3
    doc.add_paragraph()
    h2 = doc.add_heading("3.3  The Database Class — db/database.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Instead of writing raw SQL everywhere in the code, all database interactions "
        "are wrapped inside a single class called JobDatabase. This is a common "
        "software pattern called the Repository Pattern — it keeps database logic "
        "in one place."
    )

    doc.add_paragraph()
    doc.add_paragraph("How the class initializes itself:").runs[0].font.bold = True
    add_code_block(
        doc,
        "class JobDatabase:\n"
        "    def __init__(self):\n"
        "        current_dir = Path(__file__).parent\n"
        "        self.db_path = current_dir / 'jobs.sqlite'  # the database file\n"
        "        self.schema_path = current_dir / 'schema.sql'\n"
        "        self._init_db()  # create the table if it doesn't exist yet\n"
        "\n"
        "    def _init_db(self):\n"
        "        with open(self.schema_path) as f:\n"
        "            schema = f.read()\n"
        "        with sqlite3.connect(self.db_path) as conn:\n"
        "            conn.executescript(schema)  # run the CREATE TABLE statement",
    )
    add_info_box(
        doc,
        "The CREATE TABLE IF NOT EXISTS in schema.sql means this is safe to run every "
        "time the app starts. If the table already exists, nothing happens.",
    )

    doc.add_paragraph()
    doc.add_paragraph("Key methods:").runs[0].font.bold = True
    methods = [
        ("add_job(job_data)",
         "Inserts a new job record. The requirements and benefits lists are converted "
         "to JSON strings with json.dumps() before saving."),
        ("get_all_jobs()",
         "Returns every job as a list of Python dictionaries. JSON strings are "
         "converted back to lists with json.loads()."),
        ("search_jobs(skills, experience_level)",
         "The most important method for the AI pipeline. It builds a SQL query "
         "dynamically: for each skill the candidate has, it adds a LIKE condition "
         "to search inside the requirements column."),
    ]
    for method, desc in methods:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(method).font.bold = True
        p.add_run(": " + desc)

    # 3.4
    doc.add_paragraph()
    h2 = doc.add_heading("3.4  Seeding the Database — db/seed_jobs.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "A fresh database is empty. The seed_jobs.py script populates it with sample "
        "job listings so the matcher has something to work with. It defines a list of "
        "job dictionaries and calls db.add_job() for each one."
    )
    add_code_block(
        doc,
        "# Run once to populate the database:\n"
        "python db/seed_jobs.py\n"
        "\n"
        "# Sample job entry:\n"
        "{\n"
        '    "title": "Senior Software Engineer",\n'
        '    "company": "TechCorp",\n'
        '    "location": "Remote",\n'
        '    "type": "Full-time",\n'
        '    "experience_level": "Senior",\n'
        '    "salary_range": "$120,000 - $180,000",\n'
        '    "requirements": ["Python", "JavaScript", "React", "AWS", "Kubernetes"],\n'
        '    "benefits": ["Health insurance", "401(k) matching", "Remote work"]\n'
        "}",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. The Agent Architecture
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("4. The Agent Architecture", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "An 'agent' in this project is simply a Python class that has a specific job "
        "to do. Instead of one huge function that does everything, the work is split "
        "among specialist agents — just like different departments in a real company."
    )
    add_info_box(
        doc,
        "Think of the agents like an assembly line: each agent receives the output of "
        "the previous one, adds its own work, and passes everything forward.",
    )

    # 4.1
    doc.add_paragraph()
    h2 = doc.add_heading("4.1  The Base Agent — agents/base_agent.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "All agents share the same foundation: the BaseAgent class. It handles the "
        "connection to Ollama and provides two helper methods that every agent uses."
    )
    add_code_block(
        doc,
        "class BaseAgent:\n"
        "    def __init__(self, name: str, instructions: str):\n"
        "        self.name = name\n"
        "        self.instructions = instructions  # the agent's 'system prompt'\n"
        "        self.ollama_client = OpenAI(\n"
        '            base_url="http://localhost:11434/v1",  # local Ollama server\n'
        '            api_key="ollama",  # required by the library but not checked\n'
        "        )",
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "The two helper methods every agent inherits:"
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("_query_ollama(prompt)").font.bold = True
    p.add_run(
        ": Sends the agent's instructions plus a user prompt to the Llama 3.2 model "
        "and returns the model's text response. The instructions act as the agent's "
        "personality — they tell the model what role it is playing."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("_parse_json_safely(text)").font.bold = True
    p.add_run(
        ": AI models sometimes wrap JSON in extra text. This method searches the "
        "response for the first { and last } to extract only the JSON part, then "
        "parses it. If parsing fails it returns an error dict instead of crashing."
    )

    # 4.2
    doc.add_paragraph()
    h2 = doc.add_heading("4.2  Extractor Agent — agents/extractor_agent.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph("Responsibility: Convert the PDF resume into structured data.")
    doc.add_paragraph()
    doc.add_paragraph("Step-by-step what it does:").runs[0].font.bold = True
    steps = [
        "Receives the file path of the uploaded PDF.",
        "Uses pdfminer.six's extract_text() to pull all readable text from the PDF.",
        "Sends that raw text to Ollama with the instruction to 'extract and structure "
        "information from resumes — focus on personal info, work experience, education, "
        "skills, and certifications'.",
        "Returns a dictionary containing the raw_text, the structured_data from the "
        "AI, and an extraction_status field.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        p.paragraph_format.left_indent = Inches(0.3)

    add_code_block(
        doc,
        "# Simplified view of extractor_agent.py\n"
        "class ExtractorAgent(BaseAgent):\n"
        "    def __init__(self):\n"
        '        super().__init__(name="Extractor", instructions="...")\n'
        "\n"
        "    async def run(self, messages):\n"
        "        resume_data = eval(messages[-1]['content'])\n"
        '        raw_text = extract_text(resume_data["file_path"])  # PDF → text\n'
        "        extracted_info = self._query_ollama(raw_text)      # text → AI\n"
        "        return {\n"
        '            "raw_text": raw_text,\n'
        '            "structured_data": extracted_info,\n'
        '            "extraction_status": "completed"\n'
        "        }",
    )

    # 4.3
    doc.add_paragraph()
    h2 = doc.add_heading("4.3  Analyzer Agent — agents/analyzer_agent.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Responsibility: Turn the unstructured AI text into a reliable, machine-readable "
        "skills profile."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "This is the most complex agent. The challenge is that an AI model does not "
        "always return perfectly formatted JSON — it sometimes adds extra words or "
        "uses the wrong format. The AnalyzerAgent handles this with two strategies:"
    )
    p = doc.add_paragraph(style="List Number")
    p.add_run("Primary: LLM extraction. ").font.bold = True
    p.add_run(
        "It gives the AI a very specific prompt that says 'Return ONLY valid JSON' "
        "and specifies the exact structure expected."
    )
    p = doc.add_paragraph(style="List Number")
    p.add_run("Fallback: Regex extraction. ").font.bold = True
    p.add_run(
        "If the AI returns broken JSON or an empty skills list, the agent has a "
        "_extract_skills_regex() method. It contains a hardcoded list of ~60 common "
        "tech skills (python, react, docker, etc.) and simply checks whether any of "
        "them appear in the resume text. Simple but very reliable."
    )

    add_code_block(
        doc,
        "# The prompt sent to the AI (simplified)\n"
        '"""\n'
        "Analyze this resume data and extract key information.\n"
        "IMPORTANT: Return ONLY valid JSON, no explanations.\n"
        "\n"
        "Return this exact structure:\n"
        "{\n"
        '    "technical_skills": ["python", "javascript"],\n'
        '    "years_of_experience": 5,\n'
        '    "education": {"level": "Bachelors", "field": "CS"},\n'
        '    "experience_level": "Mid-level",\n'
        '    "key_achievements": ["..."],\n'
        '    "domain_expertise": ["..."]\n'
        "}\n"
        '"""',
    )
    doc.add_paragraph()
    doc.add_paragraph("What it returns to the next agent:")
    add_code_block(
        doc,
        "{\n"
        '    "skills_analysis": {\n'
        '        "technical_skills": ["python", "sql", "machine learning"],\n'
        '        "years_of_experience": 4,\n'
        '        "education": {"level": "Bachelors", "field": "Computer Science"},\n'
        '        "experience_level": "Mid-level",\n'
        '        "key_achievements": ["Built ML pipeline", "Led team of 5"],\n'
        '        "domain_expertise": ["Data Science", "Backend"]\n'
        "    },\n"
        '    "analysis_timestamp": "2026-04-23T10:00:00",\n'
        '    "confidence_score": 0.85\n'
        "}",
    )

    # 4.4
    doc.add_paragraph()
    h2 = doc.add_heading("4.4  Matcher Agent — agents/matcher_agent.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Responsibility: Find job listings from the database that best fit the candidate."
    )
    doc.add_paragraph()
    doc.add_paragraph("This agent bridges the AI world and the database world:").runs[
        0
    ].font.bold = True
    steps = [
        "Reads the skills list and experience_level from the analyzer's output.",
        "Calls search_jobs() which queries the SQLite database using LIKE conditions — "
        "for example: WHERE requirements LIKE '%python%' OR requirements LIKE '%sql%'.",
        "For each job found, calculates a match_score by comparing the set of the "
        "candidate's skills with the set of required skills for that job.",
        "Only keeps jobs where the match score is at least 50% OR at least one skill "
        "matches exactly.",
        "Returns the list sorted from highest to lowest match score.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        p.paragraph_format.left_indent = Inches(0.3)

    add_code_block(
        doc,
        "# How the match score is calculated\n"
        "required_skills = set(job['requirements'])       # e.g. {python, react, aws}\n"
        "candidate_skills = set(skills)                   # e.g. {python, sql, docker}\n"
        "\n"
        "overlap = len(required_skills & candidate_skills)  # intersection = {python}\n"
        "match_score = int((overlap / len(required_skills)) * 100)  # = 33%\n"
        "\n"
        "# Job is included if score >= 50% OR at least 1 skill matches\n"
        "if match_score >= 50 or overlap >= 1:\n"
        "    scored_jobs.append({...})",
    )

    # 4.5
    doc.add_paragraph()
    h2 = doc.add_heading("4.5  Screener Agent — agents/screener_agent.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Responsibility: Produce a human-readable screening report about the candidate."
    )
    doc.add_paragraph(
        "This agent receives the entire workflow context (everything collected so far) "
        "and passes it to Ollama with the instruction to evaluate: qualification "
        "alignment, experience relevance, skill match percentage, cultural fit indicators, "
        "and any red flags. It returns a free-text screening_report along with a "
        "hardcoded screening_score of 85 (which in a real system would be calculated "
        "from the matched jobs)."
    )

    # 4.6
    doc.add_paragraph()
    h2 = doc.add_heading("4.6  Recommender Agent — agents/recommender_agent.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Responsibility: Generate the final human-readable recommendation."
    )
    doc.add_paragraph(
        "This is the last agent in the chain. It receives the full workflow context "
        "and asks Ollama to produce a final recommendation considering all previous "
        "stages: the extracted profile, the skills analysis, the job matches, and the "
        "screening results. The output is a plain-text paragraph with clear next steps "
        "for the recruiter."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. The Orchestrator
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("5. The Orchestrator — agents/orchestrator.py", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "The Orchestrator is the conductor of the orchestra. It does not process "
        "resumes directly — it creates all the other agents and calls them in the "
        "correct order, passing data between them."
    )

    doc.add_paragraph()
    doc.add_paragraph("The complete pipeline in code:").runs[0].font.bold = True
    add_code_block(
        doc,
        "async def process_application(self, resume_data):\n"
        "    # Step 1: Extract text and structure from the PDF\n"
        "    extracted_data = await self.extractor.run([...])\n"
        "\n"
        "    # Step 2: Analyze skills, experience, education\n"
        "    analysis_results = await self.analyzer.run([extracted_data])\n"
        "\n"
        "    # Step 3: Search the database for matching jobs\n"
        "    job_matches = await self.matcher.run([analysis_results])\n"
        "\n"
        "    # Step 4: Screen the candidate holistically\n"
        "    screening_results = await self.screener.run([full_context])\n"
        "\n"
        "    # Step 5: Generate the final recommendation\n"
        "    final_recommendation = await self.recommender.run([full_context])\n"
        "\n"
        "    return workflow_context  # the complete result",
    )

    doc.add_paragraph()
    add_info_box(
        doc,
        "Notice the async/await keywords. Python's asyncio is used so that the app "
        "stays responsive while waiting for the AI model to respond — which can take "
        "several seconds per call.",
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The orchestrator also maintains a workflow_context dictionary that grows "
        "at each step:"
    )
    add_code_block(
        doc,
        "workflow_context = {\n"
        '    "resume_data":          {...},   # the uploaded file path\n'
        '    "extracted_data":       {...},   # added after step 1\n'
        '    "analysis_results":     {...},   # added after step 2\n'
        '    "job_matches":          {...},   # added after step 3\n'
        '    "screening_results":    {...},   # added after step 4\n'
        '    "final_recommendation": {...},   # added after step 5\n'
        '    "status":               "completed"\n'
        "}",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. The Web Interface
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("6. The Web Interface — app.py", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "app.py is the entry point of the application. It builds the entire web "
        "interface using Streamlit — a Python library that turns scripts into "
        "interactive web pages."
    )

    h2 = doc.add_heading("6.1  How the UI Is Built", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Streamlit works by running the Python script from top to bottom every time "
        "the user interacts with the page. Key UI components used:"
    )
    bullets = [
        ("st.sidebar + option_menu",
         "Creates the left navigation panel with menu items."),
        ("st.file_uploader(type=['pdf'])",
         "Provides the file upload button restricted to PDF files."),
        ("st.spinner / st.progress",
         "Shows a loading indicator and progress bar while processing."),
        ("st.tabs([...])",
         "Splits the results into four tabs: Analysis, Job Matches, Screening, Recommendation."),
        ("st.metric",
         "Displays the confidence score and screening score as highlighted numbers."),
    ]
    for widget, desc in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(widget).font.bold = True
        p.add_run(": " + desc)

    doc.add_paragraph()
    h2 = doc.add_heading("6.2  Connecting UI to Agents", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "When the user uploads a PDF, the app follows these steps:"
    )
    steps = [
        "save_uploaded_file() saves the PDF to the uploads/ folder with a timestamp "
        "in the filename (e.g. uploads/resume_20260423_103000_cv.pdf).",
        "asyncio.run(process_resume(file_path)) starts the async pipeline. "
        "asyncio.run() bridges the synchronous Streamlit world with the async agents.",
        "process_resume() creates an OrchestratorAgent and calls process_application().",
        "When the result comes back, the app reads the workflow_context dictionary "
        "and displays different parts in the appropriate tabs.",
        "The full result is also saved as a text file in the results/ folder.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. Utilities
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("7. Utilities", 1)
    set_heading_color(h1, 0, 102, 204)

    h2 = doc.add_heading("7.1  Logger — utils/logger.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Every time the app runs, setup_logger() creates a new log file in the logs/ "
        "folder with a timestamp in the name (e.g. recruitment_20260423_103000.log). "
        "All messages are printed both to the file and to the terminal simultaneously, "
        "using Python's built-in logging module."
    )

    doc.add_paragraph()
    h2 = doc.add_heading("7.2  Custom Exceptions — utils/exceptions.py", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Custom exception classes make error messages more descriptive and allow "
        "specific error types to be caught independently. They all inherit from "
        "ResumeProcessingError which itself inherits from Python's built-in Exception:"
    )
    add_code_block(
        doc,
        "ResumeProcessingError  ← base for all errors\n"
        "    ├── ExtractionError    (PDF reading failed)\n"
        "    ├── AnalysisError      (skill parsing failed)\n"
        "    ├── MatchingError      (database query failed)\n"
        "    ├── ScreeningError     (screening step failed)\n"
        "    └── RecommendationError (recommendation step failed)",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. End-to-End Flow Diagram
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("8. End-to-End Flow", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "Here is the complete journey of a resume through the system, from upload to "
        "final recommendation:"
    )
    doc.add_paragraph()
    flow = (
        "  USER uploads PDF\n"
        "       │\n"
        "       ▼\n"
        "  app.py  →  save_uploaded_file()  →  uploads/resume_*.pdf\n"
        "       │\n"
        "       ▼\n"
        "  OrchestratorAgent.process_application()\n"
        "       │\n"
        "       ├─ 1 ─► ExtractorAgent\n"
        "       │          pdfminer extracts raw text\n"
        "       │          Llama 3.2 structures it\n"
        "       │          → {raw_text, structured_data}\n"
        "       │\n"
        "       ├─ 2 ─► AnalyzerAgent\n"
        "       │          Llama 3.2 extracts skills & experience\n"
        "       │          Regex fallback if AI fails\n"
        "       │          → {technical_skills, experience_level, ...}\n"
        "       │\n"
        "       ├─ 3 ─► MatcherAgent\n"
        "       │          SQL LIKE query on jobs.sqlite\n"
        "       │          Calculates % match score per job\n"
        "       │          → {matched_jobs: [{title, match_score, location}]}\n"
        "       │\n"
        "       ├─ 4 ─► ScreenerAgent\n"
        "       │          Llama 3.2 writes screening report\n"
        "       │          → {screening_report, screening_score}\n"
        "       │\n"
        "       └─ 5 ─► RecommenderAgent\n"
        "                  Llama 3.2 writes final recommendation\n"
        "                  → {final_recommendation}\n"
        "       │\n"
        "       ▼\n"
        "  app.py displays results in 4 tabs\n"
        "  Saves full result to results/analysis_*.txt"
    )
    add_code_block(doc, flow)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. How to Run
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("9. How to Run the Project", 1)
    set_heading_color(h1, 0, 102, 204)

    h2 = doc.add_heading("Step 1 — Install Ollama", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph(
        "Download and install Ollama from https://ollama.com. After installation, "
        "pull the Llama 3.2 model:"
    )
    add_code_block(doc, "ollama pull llama3.2")

    doc.add_paragraph()
    h2 = doc.add_heading("Step 2 — Install Python Dependencies", 2)
    set_heading_color(h2, 0, 128, 0)
    add_code_block(doc, "pip install -r requirements.txt")

    doc.add_paragraph()
    h2 = doc.add_heading("Step 3 — Seed the Database", 2)
    set_heading_color(h2, 0, 128, 0)
    doc.add_paragraph("This creates the SQLite database and fills it with sample jobs:")
    add_code_block(doc, "python db/seed_jobs.py")

    doc.add_paragraph()
    h2 = doc.add_heading("Step 4 — Start the App", 2)
    set_heading_color(h2, 0, 128, 0)
    add_code_block(doc, "streamlit run app.py")
    doc.add_paragraph(
        "The browser will open automatically at http://localhost:8501. Upload a PDF "
        "resume and watch the agents process it in real time."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. Summary
    # ══════════════════════════════════════════════════════════════
    h1 = doc.add_heading("10. Summary", 1)
    set_heading_color(h1, 0, 102, 204)

    doc.add_paragraph(
        "Here is a quick-reference summary of all the components and their roles:"
    )
    doc.add_paragraph()

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    for i, col in enumerate(["File / Folder", "Layer", "Responsibility"]):
        hdr[i].text = col
        hdr[i].paragraphs[0].runs[0].font.bold = True
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0066CC")
        tcPr.append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    summary_rows = [
        ("app.py", "Frontend", "Streamlit UI, file upload, display results"),
        ("db/schema.sql", "Database", "Defines the jobs table structure"),
        ("db/database.py", "Database", "JobDatabase class — all SQL operations"),
        ("db/seed_jobs.py", "Database", "Populates the DB with sample jobs"),
        ("agents/base_agent.py", "Agent Core", "Connects to Ollama, shared helpers"),
        ("agents/orchestrator.py", "Agent Core", "Runs the 5-step pipeline"),
        ("agents/extractor_agent.py", "Agent", "PDF → structured text via LLM"),
        ("agents/analyzer_agent.py", "Agent", "Text → skills profile"),
        ("agents/matcher_agent.py", "Agent", "Skills → matched jobs from DB"),
        ("agents/screener_agent.py", "Agent", "Candidate → screening report"),
        ("agents/recommender_agent.py", "Agent", "Full context → recommendation"),
        ("utils/logger.py", "Utility", "Writes logs to logs/ folder"),
        ("utils/exceptions.py", "Utility", "Custom error types"),
    ]
    for file, layer, resp in summary_rows:
        row = summary_table.add_row().cells
        row[0].text = file
        row[0].paragraphs[0].runs[0].font.name = "Courier New"
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = layer
        row[2].text = resp

    doc.add_paragraph()
    doc.add_paragraph()
    closing = doc.add_paragraph(
        "You now have a complete understanding of every layer of the AI Recruiter Agency — "
        "from the SQLite database that stores jobs, through the specialized AI agents "
        "that process resumes, to the Streamlit interface that ties it all together. "
        "Happy coding!"
    )
    closing.paragraph_format.space_before = Pt(12)

    # ── Save ──────────────────────────────────────────────────────
    output_path = "tutorial.docx"
    doc.save(output_path)
    print(f"Tutorial saved to {output_path}")


if __name__ == "__main__":
    build_tutorial()
